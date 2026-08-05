"""
Builds the Hairline Salon Manager prototype user guide as a styled Word document,
with a real screenshot of every screen.
"""

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

SHOTS = Path(
    r"C:\temp\claude\c--Data-OneDrive---Terra-Group-Applications-Hairline"
    r"\6305bd13-c82c-4c05-b471-9115eecd7529\scratchpad\guide-shots"
)
OUT = Path(r"C:\tmp\hairline-proto\docs\Hairline Salon Manager - User Guide.docx")

TAUPE = RGBColor(0x8A, 0x7F, 0x6F)
TAUPE_DEEP = RGBColor(0x6E, 0x64, 0x55)
INK = RGBColor(0x1A, 0x18, 0x16)
BODY = RGBColor(0x3A, 0x36, 0x2F)
MUTED = RGBColor(0x7A, 0x72, 0x64)
CRIT = RGBColor(0xA0, 0x43, 0x3A)
WARN = RGBColor(0xA8, 0x76, 0x2A)

FONT = "Segoe UI"
FONT_LIGHT = "Segoe UI Light"
IMG_W = Inches(6.3)

doc = Document()

# ----------------------------------------------------------------- page setup
for section in doc.sections:
    section.page_width = Inches(8.27)   # A4
    section.page_height = Inches(11.69)
    section.left_margin = Inches(0.95)
    section.right_margin = Inches(0.95)
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)

style = doc.styles["Normal"]
style.font.name = FONT
style.font.size = Pt(10.5)
style.font.color.rgb = BODY
style.paragraph_format.space_after = Pt(7)
style.paragraph_format.line_spacing = 1.28


def shade(paragraph, hex_colour):
    """Paints a paragraph background — used for callouts."""
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_colour)
    pPr.append(shd)


def border(paragraph, edge="bottom", colour="E2DED7", size=6):
    pPr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    el = OxmlElement(f"w:{edge}")
    el.set(qn("w:val"), "single")
    el.set(qn("w:sz"), str(size))
    el.set(qn("w:space"), "6")
    el.set(qn("w:color"), colour)
    borders.append(el)
    pPr.append(borders)


def text(txt, size=10.5, colour=BODY, bold=False, italic=False, font=FONT,
         align=None, space_before=0, space_after=7, spacing=1.28, caps=False):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = spacing
    run = p.add_run(txt.upper() if caps else txt)
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = colour
    return p


def rich(parts, size=10.5, space_after=7, spacing=1.28):
    """A paragraph built from (text, bold, colour) segments."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = spacing
    for seg in parts:
        txt, bold, colour = (seg + (BODY,))[:3] if len(seg) == 2 else seg
        run = p.add_run(txt)
        run.font.name = FONT
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.color.rgb = colour
    return p


def h1(txt):
    p = text(txt, size=22, colour=INK, font=FONT_LIGHT, space_before=2, space_after=4,
             spacing=1.1)
    return p


def h2(txt, eyebrow=None):
    if eyebrow:
        text(eyebrow, size=8.5, colour=TAUPE, bold=True, caps=True,
             space_before=14, space_after=2)
    p = text(txt, size=16, colour=INK, font=FONT_LIGHT, space_after=3, spacing=1.12)
    border(p)
    return p


def h3(txt):
    return text(txt, size=11.5, colour=INK, bold=True, space_before=9, space_after=3)


def bullet(txt, bold_head=None):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.22)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.26
    dash = p.add_run("—   ")
    dash.font.name = FONT
    dash.font.size = Pt(10.5)
    dash.font.color.rgb = TAUPE
    if bold_head:
        rb = p.add_run(bold_head)
        rb.font.name = FONT
        rb.font.size = Pt(10.5)
        rb.font.bold = True
        rb.font.color.rgb = INK
    r = p.add_run(txt)
    r.font.name = FONT
    r.font.size = Pt(10.5)
    r.font.color.rgb = BODY
    return p


def numbered(n, txt, bold_head=None):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.22)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.26
    num = p.add_run(f"{n}.  ")
    num.font.name = FONT
    num.font.size = Pt(10.5)
    num.font.bold = True
    num.font.color.rgb = TAUPE_DEEP
    if bold_head:
        rb = p.add_run(bold_head)
        rb.font.name = FONT
        rb.font.size = Pt(10.5)
        rb.font.bold = True
        rb.font.color.rgb = INK
    r = p.add_run(txt)
    r.font.name = FONT
    r.font.size = Pt(10.5)
    r.font.color.rgb = BODY
    return p


def callout(title, txt, fill="F1EEE8", colour=TAUPE_DEEP):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.1)
    p.paragraph_format.right_indent = Inches(0.1)
    shade(p, fill)
    rt = p.add_run(title + "  ")
    rt.font.name = FONT
    rt.font.size = Pt(10)
    rt.font.bold = True
    rt.font.color.rgb = colour
    rb = p.add_run(txt)
    rb.font.name = FONT
    rb.font.size = Pt(10)
    rb.font.color.rgb = colour
    return p


def screenshot(name, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    p.add_run().add_picture(str(SHOTS / f"{name}.png"), width=IMG_W)
    cap = text(caption, size=8.5, colour=MUTED, italic=True,
               align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
    return cap


def wordmark_para(size=30, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    for txt, colour in (("HAIR", TAUPE), ("|", INK), ("line", INK)):
        r = p.add_run(txt)
        r.font.name = FONT_LIGHT
        r.font.size = Pt(size)
        r.font.color.rgb = colour
    return p


def table(rows, widths, header=True):
    t = doc.add_table(rows=0, cols=len(widths))
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for r_i, row in enumerate(rows):
        cells = t.add_row().cells
        for c_i, value in enumerate(row):
            cell = cells[c_i]
            cell.width = widths[c_i]
            para = cell.paragraphs[0]
            para.paragraph_format.space_after = Pt(3)
            para.paragraph_format.space_before = Pt(3)
            run = para.add_run(value.upper() if (header and r_i == 0) else value)
            run.font.name = FONT
            run.font.size = Pt(8.5 if (header and r_i == 0) else 10)
            run.font.bold = (header and r_i == 0) or c_i == 0
            run.font.color.rgb = MUTED if (header and r_i == 0) else (
                INK if c_i == 0 else BODY
            )
            border(para, "bottom", "E2DED7" if (header and r_i == 0) else "EDEAE4",
                   8 if (header and r_i == 0) else 4)
    return t


def page_break():
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


# ================================================================= cover
for _ in range(3):
    doc.add_paragraph()

wordmark_para(46)
p = text("Salon Manager", size=30, colour=INK, font=FONT_LIGHT, space_after=2)
text("Prototype user guide", size=14, colour=TAUPE_DEEP, space_after=18)
border(p)

text("A walk through every screen of the working prototype, with pictures of each one, "
     "so you can see what the new system does before a line of it is signed off.",
     size=12, colour=BODY, space_after=16, spacing=1.4)

table([
    ["Prepared for", "The owner of Hairline"],
    ["Version", "Prototype v1 — August 2026"],
    ["Built on", "Hairline's own MySalon data, backup of 29 July 2026"],
    ["Demo trading day", "Saturday, 25 July 2026"],
], [Inches(1.9), Inches(4.4)], header=False)

doc.add_paragraph()
callout("Client names are anonymised.",
        "Services, prices, products, stock, revenue, staff and visit patterns are all real. "
        "Client names, phone numbers, e-mail addresses and birthdays have been replaced with "
        "realistic stand-ins so the prototype can be shared safely.")

page_break()

# ================================================================= contents
h1("What's in this guide")
text("Every screen in the prototype, in the order you will meet it.", colour=MUTED,
     space_after=12)

contents = [
    ["1", "Signing in", "Who gets in, and what each person sees"],
    ["2", "The dashboard", "The business at a glance"],
    ["3", "The till", "Ringing up a sale in under 30 seconds"],
    ["4", "Clients", "Finding a client and reading their history"],
    ["5", "The diary", "The day, stylist by stylist"],
    ["6", "Stock", "Retail, back bar and what to order"],
    ["7", "The team", "Staff portfolios and time clock"],
    ["8", "Cash-up", "Counting the drawer and locking the day"],
    ["9", "Pricing", "The menu, margins and the printed price list"],
    ["10", "On a phone", "What the stylists carry with them"],
]
table([["", "Screen", "What it covers"]] + contents,
      [Inches(0.4), Inches(1.8), Inches(4.1)])

doc.add_paragraph()
callout("Before you start.",
        "Nothing you do in the prototype can break anything. Sales you ring up are kept in "
        "your own browser only, and they disappear when you clear it. Click freely.")

page_break()

# ================================================================= 1 sign in
h2("Signing in", "Section 1")
text("The prototype opens on a sign-in screen. Four sign-ins are provided, one for each kind "
     "of person in the salon. They are listed on the screen itself, and clicking one fills in "
     "the form for you.", space_after=8)

table([
    ["Username", "Signs you in as", "What they can reach"],
    ["owner", "Salon Owner", "Everything: takings, reports, costs, stock, the team"],
    ["reception", "Reception", "Till, clients, diary, stock, cash-up"],
    ["karin", "Karin M.", "One stylist's own day, figures and tips"],
    ["meagan", "Meagan V.", "A second stylist, to compare"],
], [Inches(1.3), Inches(1.6), Inches(3.4)])

doc.add_paragraph()
rich([("The password for all four is  ", False), ("hairline2026", True, INK), (".", False)])

screenshot("01-login",
           "The sign-in screen. Click any of the four demo sign-ins on the right and the form "
           "fills itself in.")

callout("A note on this sign-in.",
        "This screen keeps the demo link away from casual visitors, but it is not real "
        "security — the prototype has no server behind it. Proper accounts, passwords and "
        "permissions are part of the production build, not the prototype.",
        fill="F7EFE0", colour=WARN)

rich([("To change who you are signed in as, press  ", False),
      ("Sign out", True, INK),
      ("  at the bottom of the menu on the left, then sign in as someone else.", False)],
     space_after=4)

page_break()

# ================================================================= 2 dashboard
h2("The dashboard", "Section 2")
text("This is what the owner sees first. It answers the question you would otherwise phone "
     "the salon to ask: how is today going, and is anything going wrong?", space_after=8)

h3("Reading the screen")
bullet("shows what has been taken so far, how many sales, and the average spend per client.",
       bold_head="Today ")
bullet("compares every full year since 2015, so a good or bad year is obvious at a glance.",
       bold_head="Revenue by year ")
bullet("shows the last two years month by month, with the best month named underneath.",
       bold_head="Monthly revenue ")
bullet("ranks the team over the last twelve months, including how much retail each one sells.",
       bold_head="Top stylists ")
bullet("holds the four numbers that need action — lapsed clients, stock needing a count, "
       "birthdays missing, and your loyal core.", bold_head="Worth your attention ")

screenshot("02-dashboard",
           "The owner's dashboard. Every figure is drawn from Hairline's own records.")

callout("The four warnings are real.",
        "1,350 stock lines currently show a negative quantity, 758 clients have not been in "
        "for over 90 days, and only 8% of clients have a birthday on file. These are not "
        "sample numbers — they are what the salon's data says today.")

page_break()

# ================================================================= 3 till
h2("The till", "Section 3")
text("The most important screen in the system. Everything about it is built around one "
     "target: a routine sale rung up in under thirty seconds.", space_after=8)

h3("Ringing up a sale")
numbered(1, "by name or phone, or press ", bold_head="Find the client ")
numbered(2, "the services from the tabs, then switch to Retail for products. Each one you "
            "click drops into the sale on the right.", bold_head="Tap ")
numbered(3, "if you need to: change the stylist, the quantity, or add a discount "
            "percentage on any line.", bold_head="Adjust the line ")
numbered(4, "if the client leaves one — it is recorded against the stylist, and kept out of "
            "the sale total.", bold_head="Add a tip ")
numbered(5, "Card, cash, EFT, voucher or account. You can split across several — the "
            "balance updates as you go.", bold_head="Take the payment. ")
numbered(6, "The button turns solid once the sale is covered.",
         bold_head="Press Complete sale. ")

screenshot("03-till-empty",
           "The till before a sale starts. Services are grouped the way the salon works, with "
           "the most-used department first.")

screenshot("05-till-payment",
           "A sale in progress: a real client, two services, VAT shown, and a card payment "
           "covering the balance. The timer in the top right reads 5 seconds.")

h3("Things worth knowing")
bullet("counts from the moment you start the sale, so you can see the 30-second target "
       "being met or missed.", bold_head="The timer ")
bullet("is shown for information, calculated at 15% inclusive — the prices are what the "
       "client pays.", bold_head="VAT ")
bullet("is offered before you pick anyone, so a walk-in never slows you down.",
       bold_head="Walk-in ")
bullet("more than the total gives change; a card or voucher can never be over-captured.",
       bold_head="Cash tendered ")

page_break()

# ================================================================= 4 clients
h2("Clients", "Section 4")
text("Eleven years of visit history, searchable in a keystroke. This is the salon's most "
     "valuable asset and the hardest thing to replace if it were ever lost.", space_after=8)

h3("Finding someone")
bullet("Type any part of a name or a phone number into the search box.")
bullet("filter the list down to lapsed clients, top spenders, or birthdays this month.",
       bold_head="The chips ")
bullet("Click any name to open their file.")

screenshot("06-clients",
           "The client list, sorted by who was in most recently. Badges flag VIPs, medical "
           "notes and lapsed clients.")

h3("Inside a client's file")
bullet("across the top: visits, lifetime spend, average visit and when they were last in.",
       bold_head="Four figures ")
bullet("every visit, every service and product, the price paid and the stylist who did it.",
       bold_head="The timeline: ")
bullet("colour formulas and preferences, which travel with the client to whoever serves them.",
       bold_head="Stylist notes: ")
bullet("opens a message with a template already filled in — useful for winning back someone "
       "who has drifted away.", bold_head="Send message ")

screenshot("07-client-file",
           "A client file. The visit timeline goes back as far as the salon's records do.")

page_break()

# ================================================================= 5 diary
h2("The diary", "Section 5")
text("A light appointment book, one column per stylist. It is deliberately simple: the diary "
     "in MySalon was barely used in eleven years, so this exists for the team to try without "
     "anyone being forced onto it.", space_after=8)

bullet("Colours show the department, so a day of colour work is obvious at a glance.")
bullet("Click any appointment for the client, the service, the stylist and what it was worth.")
bullet("Overlapping appointments sit side by side rather than hiding each other.")
bullet("Invoicing never requires a booking, so walk-ins are unaffected.")

screenshot("08-diary",
           "The demo day, reconstructed from that day's real invoices. Longer blocks are "
           "colour and extension work.")

callout("How these appointments were worked out.",
        "MySalon records the moment a client pays, not when they sat down. The prototype works "
        "backwards from checkout using a realistic time for each service, so the day reads the "
        "way it was actually worked. In the real system these are simply the bookings you made.")

page_break()

# ================================================================= 6 stock
h2("Stock", "Section 6")
text("Retail shelf and professional back bar are counted separately, which is what the "
     "industry recommends and what Hairline already does informally with its brand "
     "departments.", space_after=8)

h3("The three tabs")
bullet("everything sold to clients, with cost, selling price, margin and quantity on hand.",
       bold_head="Retail shelf: ")
bullet("the professional products used during services, valued at cost.",
       bold_head="Back bar: ")
bullet("everything at or below its reorder level, grouped by supplier, with a suggested "
       "quantity.", bold_head="What to order: ")

screenshot("09-stock",
           "The retail shelf. Rows tinted amber are low; rows tinted red need a physical count.")

screenshot("10-stock-order",
           "The order list, grouped by supplier, so a phone call to one rep covers everything.")

callout("Why so many lines say “needs count”.",
        "1,350 of the salon's 3,447 stock lines currently show a negative quantity in MySalon. "
        "That happens when sales are rung up against stock that was never booked in. The "
        "prototype shows the figures exactly as they stand rather than tidying them up. A "
        "receiving step at delivery, plus one full stock take, is what turns this into a number "
        "you can trust.", fill="F8EAE8", colour=CRIT)

page_break()

# ================================================================= 7 team
h2("The team", "Section 7")
text("A profile for everyone on the books, and a portfolio for each stylist showing what "
     "they bring in.", space_after=8)

bullet("shows turnover over twelve months, invoices, retail share and a twelve-month "
       "sparkline.", bold_head="The team grid ")
bullet("adds month-against-target, tips, staff advances, the week's clocked hours and their "
       "clients for the day.", bold_head="A stylist's portfolio ")
bullet("is set at 110% of their best month in the last year — a stretch, but a real one.",
       bold_head="The monthly target ")

screenshot("11-team",
           "The team, ranked by turnover. Assistants and reception are listed separately below.")

screenshot("12-staff-portfolio",
           "A stylist's portfolio. This is the same view they see when they sign in on their "
           "own phone.")

callout("Assistants are visible for the first time.",
        "Cynthia, Thobile, Christina and Hellen clock in every day and earn substantial tips, "
        "but MySalon bills their work under the senior stylist, so no turnover is attributed to "
        "them at all. The prototype shows them honestly as assistants rather than pretending "
        "they earned nothing.")

page_break()

# ================================================================= 8 cash-up
h2("Cash-up", "Section 8")
text("The end-of-day ritual, kept exactly as reception already knows it — count the notes and "
     "coins, let the system fill in the rest.", space_after=8)

h3("Doing the cash-up")
numbered(1, "using the plus and minus buttons, or type the number straight in. The value of "
            "each row is worked out for you.", bold_head="Count each denomination ")
numbered(2, "Card, EFT and voucher totals are filled in from the day's sales — there is "
            "nothing to add up.")
numbered(3, "against expected cash. Green means balanced; red means recount.",
         bold_head="Check the variance ")
numbered(4, "Set the float you are leaving in the drawer.")
numbered(5, "Once locked, the day is closed and the figures are fixed.",
         bold_head="Press Lock the day. ")

screenshot("13-cashup",
           "The cash-up screen before counting. Card and EFT are already filled in from the "
           "till.")

callout("This particular Saturday took no cash at all.",
        "Every one of the day's 34 sales was card or EFT — R39,190 on card and R3,600 by EFT. "
        "That is not a quirk of the demo: across the last twelve months 89% of all takings came "
        "in on card and only 5% in cash.")

page_break()

# ================================================================= 9 pricing
h2("Pricing", "Section 9")
text("One place where every price lives, and the source of the printed menu clients see.",
     space_after=8)

bullet("Services are grouped by department with duration, cost, price and margin side by side.")
bullet("previews a percentage rise before anything is committed.",
       bold_head="Schedule an increase ")
bullet("produces a clean, printable menu built from these exact prices.",
       bold_head="Print client menu ")

screenshot("14-pricing",
           "The price manager. Margins are shown wherever a service cost has been captured.")

screenshot("15-price-menu",
           "The client menu, generated from the same prices the till charges — ready to print "
           "or save as a PDF.")

callout("This replaces a yearly job.",
        "Hairline currently rebuilds the client price menu as a separate Word document every "
        "year — we found five years of them in the salon's files. Here the menu is generated "
        "from the live prices, so what is printed can never drift from what clients are "
        "actually charged.")

page_break()

# ================================================================= 10 phone
h2("On a phone", "Section 10")
text("Every screen works on a phone. For stylists that is the whole point: their figures, in "
     "their pocket, without asking anyone.", space_after=8)

bullet("The menu moves to the bottom of the screen, within thumb reach.")
bullet("Stylists see only their own day, their month against target and their tips.")
bullet("The owner gets the same dashboard, so the day's takings are always a glance away.")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(10)
p.add_run().add_picture(str(SHOTS / "17-mobile-stylist.png"), width=Inches(2.5))
text("A stylist's own view on a phone.", size=8.5, colour=MUTED, italic=True,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=14)

page_break()

# ================================================================= closing
h2("What this prototype is, and is not", "In closing")

h3("What it is")
bullet("Every screen you have seen is running on Hairline's own services, prices, stock, "
       "revenue and eleven years of visit patterns.")
bullet("The till genuinely works: the totals, VAT, discounts, split payments and change are "
       "calculated properly and covered by automated tests.")
bullet("It is a faithful picture of how the finished system would look and feel.")

h3("What it is not")
bullet("It does not save anything centrally. Sales you ring up live in your browser alone.")
bullet("Messaging, ordering and scheduled price increases show the screens but do not send, "
       "order or change anything.")
bullet("The sign-in is a demo gate, not real security.")

h2("Five decisions we would like from you", "Next")
table([
    ["Name", "Happy with “Hairline Salon Manager”, or is there a name you prefer?"],
    ["Price tiers", "Confirm the stylist levels used for pricing — senior and junior, or more?"],
    ["Card machine", "Keep the standalone card terminals, or look at integrated payments later?"],
    ["WhatsApp", "Shall we register a WhatsApp Business account for the salon number?"],
    ["Diary", "Should the team use the diary from day one, or ease into it after go-live?"],
], [Inches(1.4), Inches(4.9)], header=False)

doc.add_paragraph()
p = text("This guide describes a prototype. Nothing in it is fixed, and every part of it can "
         "change on your say-so.", size=10, colour=MUTED, italic=True, space_before=10)
border(p, "top")

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(str(OUT))
print(f"Saved {OUT}  ({OUT.stat().st_size / 1024:.0f} KB)")
