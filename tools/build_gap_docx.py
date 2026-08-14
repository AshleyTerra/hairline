"""
Builds the MySalon → Hairline gap analysis as a styled Word document, in the
same house style as the user guide and the owner review.
"""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

# Width of the "—   " bullet run at 10.5pt Segoe UI, measured off the render.
DASH_W = Pt(19.2)
BULLET_X = Pt(15.84)

OUT = Path(r"C:\tmp\hairline-proto\docs\Hairline Salon Manager - MySalon Gap Analysis.docx")

TAUPE = RGBColor(0x8A, 0x7F, 0x6F)
TAUPE_DEEP = RGBColor(0x6E, 0x64, 0x55)
INK = RGBColor(0x1A, 0x18, 0x16)
BODY = RGBColor(0x3A, 0x36, 0x2F)
MUTED = RGBColor(0x7A, 0x72, 0x64)
CRIT = RGBColor(0xA0, 0x43, 0x3A)
WARN = RGBColor(0xA8, 0x76, 0x2A)
GOOD = RGBColor(0x3F, 0x6E, 0x4C)

FONT = "Segoe UI"
FONT_LIGHT = "Segoe UI Light"

doc = Document()

# ----------------------------------------------------------------- page setup
for section in doc.sections:
    section.page_width = Inches(8.27)   # A4
    section.page_height = Inches(11.69)
    section.left_margin = Inches(0.95)
    section.right_margin = Inches(0.95)
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)

style = doc.styles["Normal"]
style.font.name = FONT
style.font.size = Pt(10.5)
style.font.color.rgb = BODY
style.paragraph_format.space_after = Pt(7)
style.paragraph_format.line_spacing = 1.28


def shade(paragraph, hex_colour):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_colour)
    pPr.append(shd)


def border(paragraph, edge="bottom", colour="E2DED7", size=6):
    pPr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    el = OxmlElement(f"w:{edge}")
    el.set(qn("w:val"), "single")
    el.set(qn("w:sz"), str(size))
    el.set(qn("w:space"), "6")
    el.set(qn("w:color"), colour)
    borders.append(el)
    pPr.append(borders)


def text(txt, size=10.5, colour=BODY, bold=False, italic=False, font=FONT,
         align=None, space_before=0, space_after=7, spacing=1.28, caps=False,
         keep=False):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = spacing
    p.paragraph_format.keep_with_next = keep
    run = p.add_run(txt.upper() if caps else txt)
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = colour
    return p


def rich(parts, size=10.5, space_after=7, spacing=1.28, indent=None):
    """A paragraph built from (text, bold) or (text, bold, colour) segments."""
    p = doc.add_paragraph()
    if indent is not None:
        p.paragraph_format.left_indent = indent
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = spacing
    for seg in parts:
        txt, bold, colour = (seg + (BODY,))[:3] if len(seg) == 2 else seg
        run = p.add_run(txt)
        run.font.name = FONT
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.color.rgb = colour
    return p


def h1(txt):
    return text(txt, size=22, colour=INK, font=FONT_LIGHT, space_before=2, space_after=4,
                spacing=1.1)


def h2(txt, eyebrow=None):
    if eyebrow:
        e = text(eyebrow, size=8.5, colour=TAUPE, bold=True, caps=True,
                 space_before=14, space_after=2)
        e.paragraph_format.keep_with_next = True
    p = text(txt, size=16, colour=INK, font=FONT_LIGHT, space_after=3, spacing=1.12)
    p.paragraph_format.keep_with_next = True
    border(p)
    return p


def h3(txt):
    p = text(txt, size=11.5, colour=INK, bold=True, space_before=9, space_after=3)
    p.paragraph_format.keep_with_next = True
    return p


def area(name, verdict, verdict_colour):
    """A heading that carries its own verdict, so the page scans on the right."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(11)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(name)
    r.font.name = FONT
    r.font.size = Pt(11.5)
    r.font.bold = True
    r.font.color.rgb = INK
    sep = p.add_run("   ·   ")
    sep.font.name = FONT
    sep.font.size = Pt(11.5)
    sep.font.color.rgb = RGBColor(0xC9, 0xC3, 0xB8)
    v = p.add_run(verdict.upper())
    v.font.name = FONT
    v.font.size = Pt(8.5)
    v.font.bold = True
    v.font.color.rgb = verdict_colour
    return p


def gapline(label, items):
    rich([(label + "  ", True, MUTED), (items, False, MUTED)],
         size=9.5, space_after=4, spacing=1.24, indent=Inches(0.02))


def bullet(txt, bold_head=None):
    p = doc.add_paragraph()
    # Hanging, so a wrapped line sits under the text rather than under the dash.
    p.paragraph_format.left_indent = Pt(BULLET_X.pt + DASH_W.pt)
    p.paragraph_format.first_line_indent = Pt(-DASH_W.pt)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.26
    dash = p.add_run("—   ")
    dash.font.name = FONT
    dash.font.size = Pt(10.5)
    dash.font.color.rgb = TAUPE
    if bold_head:
        rb = p.add_run(bold_head)
        rb.font.name = FONT
        rb.font.size = Pt(10.5)
        rb.font.bold = True
        rb.font.color.rgb = INK
    r = p.add_run(txt)
    r.font.name = FONT
    r.font.size = Pt(10.5)
    r.font.color.rgb = BODY
    return p


def task(what, why):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.22)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.keep_with_next = True
    dash = p.add_run("—   ")
    dash.font.name = FONT
    dash.font.size = Pt(10.5)
    dash.font.color.rgb = TAUPE
    r = p.add_run(what)
    r.font.name = FONT
    r.font.size = Pt(10.5)
    r.font.bold = True
    r.font.color.rgb = INK
    q = doc.add_paragraph()
    q.paragraph_format.left_indent = Pt(BULLET_X.pt + DASH_W.pt)
    q.paragraph_format.space_after = Pt(5)
    q.paragraph_format.line_spacing = 1.26
    rq = q.add_run(why)
    rq.font.name = FONT
    rq.font.size = Pt(10)
    rq.font.color.rgb = BODY


def callout(title, txt, fill="F1EEE8", colour=TAUPE_DEEP):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.1)
    p.paragraph_format.right_indent = Inches(0.1)
    shade(p, fill)
    rt = p.add_run(title + "  ")
    rt.font.name = FONT
    rt.font.size = Pt(10)
    rt.font.bold = True
    rt.font.color.rgb = colour
    rb = p.add_run(txt)
    rb.font.name = FONT
    rb.font.size = Pt(10)
    rb.font.color.rgb = colour
    return p


def wordmark_para(size=30, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    for txt, colour in (("HAIR", TAUPE), ("|", INK), ("line", INK)):
        r = p.add_run(txt)
        r.font.name = FONT_LIGHT
        r.font.size = Pt(size)
        r.font.color.rgb = colour
    return p


def table(rows, widths, header=True, colours=None):
    """colours: optional per-row RGBColor applied to the last column."""
    t = doc.add_table(rows=0, cols=len(widths))
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for r_i, row in enumerate(rows):
        cells = t.add_row().cells
        is_head = header and r_i == 0
        for c_i, value in enumerate(row):
            cell = cells[c_i]
            cell.width = widths[c_i]
            para = cell.paragraphs[0]
            para.paragraph_format.space_after = Pt(3)
            para.paragraph_format.space_before = Pt(3)
            run = para.add_run(value.upper() if is_head else value)
            run.font.name = FONT
            run.font.size = Pt(8.5 if is_head else 10)
            run.font.bold = is_head or c_i == 0
            if is_head:
                run.font.color.rgb = MUTED
            elif colours and c_i == len(widths) - 1 and colours[r_i] is not None:
                run.font.color.rgb = colours[r_i]
                run.font.bold = True
                run.font.size = Pt(9)
            else:
                run.font.color.rgb = INK if c_i == 0 else BODY
            border(para, "bottom", "E2DED7" if is_head else "EDEAE4", 8 if is_head else 4)
    return t


def stat(figure, of, caption, verdict, colour=TAUPE_DEEP):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(figure)
    r.font.name = FONT_LIGHT
    r.font.size = Pt(24)
    r.font.color.rgb = INK
    o = p.add_run("   " + of)
    o.font.name = FONT
    o.font.size = Pt(10.5)
    o.font.color.rgb = MUTED
    text(caption, size=10.5, space_after=3, spacing=1.3, keep=True)
    q = text(verdict, size=10, colour=colour, bold=True, space_after=10, spacing=1.25)
    border(q, "bottom", "EDEAE4", 4)


def page_break():
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


# =================================================================== cover
for _ in range(3):
    doc.add_paragraph()

wordmark_para(46)
p = text("Salon Manager", size=30, colour=INK, font=FONT_LIGHT, space_after=2)
text("MySalon gap analysis", size=14, colour=TAUPE_DEEP, space_after=18)
border(p)

text("A feature-by-feature read of the five MySalon manuals against the working prototype — "
     "then a second pass that asks which of the missing features Hairline was ever really "
     "using.",
     size=12, colour=BODY, space_after=16, spacing=1.4)

table([
    ["Prepared for", "The owner of Hairline"],
    ["Compared", "48 pages of MySalon v8.11 manuals against 10 built screens"],
    ["Evidence", "Eleven years of migrated sales, stock and client records"],
    ["Date", "14 August 2026"],
], [Inches(1.6), Inches(4.7)], header=False)

doc.add_paragraph()
callout("The short version.",
        "Read only the manuals and you would conclude the prototype is roughly half a system. "
        "The migrated data tells a different story: several of the features it is “missing” were "
        "barely used at Hairline, and one of the biggest — stock control — is already producing "
        "numbers the salon cannot trust. The gap that matters is not the gap in the manual.")

page_break()

# =================================================================== contents
h1("What's in this document")
text("Four sections: the evidence, the coverage, the credit due to the new system, and the "
     "recommendation.", colour=MUTED, space_after=12)

table([
    ["", "Section", "What it covers"],
    ["1", "Three numbers", "What eleven years of records say about which features earned their place"],
    ["2", "Coverage", "Ten functional areas, and what is genuinely absent from each"],
    ["3", "Ahead", "Where the prototype already does more than MySalon did"],
    ["4", "Recommendation", "What to build, in order, and what to leave alone"],
], [Inches(0.4), Inches(1.6), Inches(4.3)])

doc.add_paragraph()
callout("How this was checked.",
        "Every claim about MySalon comes from the manuals in the salon's own files, read page by "
        "page. Every claim about the prototype was checked against the working system rather than "
        "from memory. The figures come from the migrated database.")

page_break()

# =================================================================== 1 numbers
h2("Three numbers that reorder the list", "Section 1")
text("A gap analysis built only from the manuals would push us to build the wrong things. These "
     "three figures are the reason.", space_after=4)

stat("R1 920", "of R7,47 million",
     "Everything ever put on account, across the whole client file. MySalon devotes four screens "
     "to account trading — payments, statements, an owing column and an account-clients filter.",
     "Don't rebuild it. Three hundredths of one percent.", CRIT)

stat("1 350", "of 3 447 stock lines",
     "Stock lines sitting at a negative quantity on hand. Another 919 sit at zero. Only 570 "
     "lines — one in six — hold a figure that could plausibly be right.",
     "Porting the stock machinery would faithfully reproduce a number nobody believes.", CRIT)

stat("76", "of 8 645 clients",
     "Client records carrying an e-mail address. Phone numbers: 8 603. Birthdays: 718. And half "
     "the book — 4 246 people — came once and never came back.",
     "SMS is the only marketing channel the data supports, and it isn't built.", CRIT)

callout("A caveat on these numbers.",
        "They describe how MySalon was used at Hairline, not whether a feature is worth having. "
        "Account trading being near zero may mean the salon chose not to offer it rather than "
        "that nobody wanted it — the same is true of loyalty. Where that judgement is yours, it "
        "is flagged as a decision rather than settled here.")

page_break()

# =================================================================== 2 coverage
h2("Coverage, area by area", "Section 2")
text("Ten functional areas, drawn from MySalon's own three navigation panels. The verdict beside "
     "each name is the summary; the line beneath names what is genuinely absent rather than "
     "merely renamed.", space_after=8)

verdicts = [
    ["Area", "Verdict"],
    ["Till and invoicing", "Ahead"],
    ["Stock and suppliers", "Largest gap"],
    ["Clients", "Partial"],
    ["Diary and appointments", "Partial"],
    ["Staff", "Partial"],
    ["Cash-up and end of day", "Partial"],
    ["Reports", "4 of about 30"],
    ["Vouchers", "At parity"],
    ["Messaging and marketing", "Not built"],
    ["Setup and administration", "Different shape"],
]
table(verdicts, [Inches(3.4), Inches(2.9)],
      colours=[None, GOOD, CRIT, WARN, WARN, WARN, WARN, WARN, GOOD, CRIT, WARN])

page_break()

area("Till and invoicing", "Ahead", GOOD)
text("Everything MySalon's invoicing panel did, plus several things it didn't: more than one "
     "docket open at once, split payment across methods on a single sale, change due on both "
     "screen and slip, and sales parked as awaiting payment.", space_after=3)
gapline("Absent:", "client quotations · account payments")

area("Stock and suppliers", "Largest gap", CRIT)
text("The prototype can look at stock, suggest what to reorder, add lines and import a supplier "
     "list. It cannot yet change what the system believes is on the shelf, except by selling it. "
     "MySalon's own rule was that stock moves three ways — received, sold, counted — and two of "
     "those three are missing.", space_after=3)
gapline("Absent:",
        "receiving against a supplier invoice · purchase orders · stock take · returns to "
        "supplier · retail-to-back-bar transfer · supplier payments and ageing · value on hand "
        "at average cost · gross profit and margin · forward cover · per-ml colour usage")

area("Clients", "Partial", WARN)
text("Client files, visit history, spend, notes and lifetime value all carry over. What is thin "
     "is everything the stylists themselves wrote down — MySalon kept a dated preference card "
     "per client, which is where the colour formulae lived.", space_after=3)
gapline("Absent:",
        "colour formula and preference history · client source · client groups · gender, area "
        "and address · loyalty points and barcode cards · default discount · photographs · "
        "statements · printable blank client card")

area("Diary and appointments", "Partial", WARN)
text("Booking works end to end: pick a date and stylist, fifteen-minute slots, clashes refused, "
     "and the booking carries through to the till as a docket — with a cancellation charge if it "
     "comes to that. What is missing sits around it.", space_after=3)
gapline("Absent:",
        "staff rota and leave planning · department colour-coding · service packages with "
        "automatic timings · waiting list · confirmation messages")

area("Staff", "Partial", WARN)
text("Staff can be added and edited, made active or inactive, and each has a portfolio with "
     "turnover, tips and retail share over any period. Monthly targets are shown but come from "
     "the migrated file — there is no way to set one, so anyone hired from here on has a target "
     "of zero.", space_after=3)
gapline("Absent:",
        "target setting · leave days and balances · employment, banking and statutory fields · "
        "notes and performance record · clocking in and out, with reasons · time-sheet report")

area("Cash-up and end of day", "Partial", WARN)
text("The drawer is counted note by note against expected takings by payment type, with the "
     "float carried forward and the variance shown. But money also leaves the drawer during the "
     "day, and there is nowhere to record that — so a day with a single petty-cash purchase "
     "cannot balance.", space_after=3)
gapline("Absent:",
        "expenditures and expenditure types · staff advances taken during the day · the flag "
        "for an advance that shouldn't hit the cash-up")

area("Reports", "4 of about 30", WARN)
text("The four that exist — staff turnover, daily staff turnover, item tracking and vouchers — "
     "are stronger than their MySalon equivalents: any date range, any combination of staff, "
     "departments and items, VAT split both ways, and every one of them exports to PDF, Excel "
     "and CSV. The count is the problem, not the quality.", space_after=3)
gapline("Absent and worth having:",
        "daily takings · business KPI (client source, gender split, average docket value) · "
        "client source · birthdays and loyalty anniversaries · purchase analysis · stock "
        "discrepancy, movement and value on hand · staff advances, tips and time sheets")
gapline("Absent, probably not missed:",
        "income statement · bank reconciliation · vendor ageing · branch transfers")

area("Vouchers", "At parity", GOOD)
text("Issued on a sale under a client or walk-in, with recipient, contact, amount, twelve-month "
     "expiry and barcode; kept off the stylist's figures as a salon sale; found by barcode, "
     "number, recipient or client; drawn down by any amount with the balance held for next time; "
     "and reported with the outstanding balance per voucher.", space_after=3)
gapline("Absent:", "vouchers about to expire shown on the dashboard")

area("Messaging and marketing", "Not built", CRIT)
text("MySalon had a whole panel: appointment confirmations, batch campaigns, templates, an inbox "
     "for replies, an audit trail and a credit balance. The prototype has a message dialog on "
     "the client file that is explicitly a preview and sends nothing.", space_after=3)
gapline("Absent:", "the entire function")

area("Setup and administration", "Different shape", WARN)
text("Users, roles and per-screen permissions are handled better than MySalon's security levels, "
     "and the prototype adds CSV import with row-by-row error reporting plus backup validation. "
     "What is missing is the configuration MySalon expected you to set up on day one.",
     space_after=3)
gapline("Absent:",
        "company details, VAT and registration numbers, logo and invoice footer · departments as "
        "numbered, named, colour-coded records · the category lists (client groups and sources, "
        "expenditure types, clock-out reasons, to-do types, resources) · invoice and stock audit "
        "trails · locking a completed invoice")

page_break()

# =================================================================== 3 ahead
h2("Where the prototype is already ahead", "Section 3")
text("Worth stating plainly, so the comparison reads fairly in both directions. None of these "
     "existed in the system being replaced.", space_after=8)

bullet("— each with their own docket, and none of them lost.",
       bold_head="Several clients at the counter at once ")
bullet("across cash, card, EFT and voucher on one sale, labelled honestly at each step.",
       bold_head="Split payment ")
bullet("shown on screen and printed on the slip.", bold_head="Change due ")
bullet("as awaiting payment, and settled later.", bold_head="Sales parked ")
bullet("— invoice, tri-fold price menu and landscape reports.",
       bold_head="Print output that fits the page ")
bullet("to PDF, Excel and CSV.", bold_head="Every report exports ")
bullet("safe across screens added in future.", bold_head="Permissions per screen, per role, ")
bullet("with a scheduled, salon-wide increase.", bold_head="A price menu builder ")
bullet("naming the row and the reason it failed.", bold_head="Imports that explain themselves, ")
bullet("— no SQL Server 2005, and backups verified before they are trusted.",
       bold_head="Runs in a browser ")

page_break()

# =================================================================== 4 plan
h2("What to build, in order", "Section 4")
text("Ordered by what the salon cannot open without, rather than by what the manual happens to "
     "document. Each tier assumes the one above it is done.", space_after=8)

h3("Tier 1 — before go-live")
text("Without these, a normal trading day cannot be recorded truthfully.",
     colour=MUTED, space_after=2, keep=True)
task("A stock take, then receiving at delivery",
     "In that order. Counting resets the ledger to something true; receiving keeps it that way. "
     "Rebuilding supplier invoicing on top of 1 350 negative lines just automates the error.")
task("Expenditures and staff advances at the cash-up",
     "Both take cash out of the drawer during the day. Until they can be entered, the variance "
     "figure is measuring the wrong thing.")
task("The client preference card, with colour formulae",
     "This is the stylists' working memory, and the one piece of client data that cannot be "
     "reconstructed. It needs somewhere to live before the old system is switched off.")
task("Departments as real records",
     "Number, name and diary colour. Every report in MySalon hangs off this, and the prototype "
     "currently treats department as a loose label on each item.")

h3("Tier 2 — the first month")
text("Not blocking, but each will be asked for within a few weeks of daily use.",
     colour=MUTED, space_after=2, keep=True)
task("Daily takings, and a KPI report",
     "The two the owner actually reads — one every morning, one at month-end. The KPI needs "
     "client source and gender captured on the client record first.")
task("Setting staff targets",
     "Monthly figure, working days, daily target and a projection. Currently display-only, which "
     "means every new hire shows a target of zero.")
task("SMS: appointment confirmations first, campaigns second",
     "99.5% of the client file has a phone number and under 1% has an e-mail. With 4 246 clients "
     "who came once, the win-back message is the highest-value thing on this whole list.")
task("Staff leave and a rota",
     "Leave balances on the staff record, and days off showing in the diary so bookings are not "
     "taken against someone who is not in.")
task("Expiring vouchers on the dashboard",
     "The machinery is built; it just needs surfacing where reception will see it.")

h3("Tier 3 — on request only")
text("Present in MySalon, and deliberately not proposed. Each is either unused at Hairline, "
     "dependent on hardware the salon does not have, or built for a multi-branch group.",
     colour=MUTED, space_after=2, keep=True)
task("Client accounts and statements",
     "R1 920 in eleven years. Four screens' worth of work for a rounding error.")
task("Loyalty points",
     "MySalon's scheme depends on barcoded loyalty cards, and no client in the file has one. "
     "Worth revisiting as a decision about the business, not as a migration task.")
task("Quotations, service packages, resource booking, phone book",
     "No trace of use in the migrated data.")
task("Income statement, bank reconciliation, forward cover, supplier ageing",
     "Accounting functions that the salon's bookkeeping already covers elsewhere.")
task("Per-ml colour control, fingerprint clock-in, branch transfers",
     "Heavyweight stock control, hardware the salon does not own, and a feature for salon groups.")

page_break()

# =================================================================== method
h2("Method, and what was read", "Notes")

text("Every feature claim about MySalon comes from the five manuals in the salon's files, read "
     "page by page. The files carry no searchable text, so all 48 pages were read as images.",
     space_after=6)

table([
    ["Manual", "Pages", "What it covered"],
    ["MS Getting Started", "9", "Set-up: staff, departments, client sources, company info, options"],
    ["MS Navigating Panels", "9", "The full menu tree — General, Reports and Manage"],
    ["MS Clients Overview", "7", "The client record card and the client reports"],
    ["MS Staff", "8", "Staff records, subs, targets and the KPI report"],
    ["MS Stock & Price List", "15", "Price list, vendor invoicing, ordering, stock take, valuation"],
], [Inches(1.9), Inches(0.6), Inches(3.8)])

doc.add_paragraph()

h3("Data handling")
bullet("Client identities in the prototype are pseudonymised, and no raw personal data has been "
       "committed to the code repository.")
bullet("The MySalon backup is used for validation only — it is checked, never restored, and it "
       "never leaves the machine it is checked on.")
bullet("Services, prices, products, stock, revenue, staff and visit patterns are all real.")

doc.add_paragraph()
p = text("This analysis describes a prototype and a proposal. Nothing in it is fixed, and every "
         "part of it can change on your say-so.",
         size=10, colour=MUTED, italic=True, space_before=10)
border(p, "top")

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(str(OUT))
print(f"Saved {OUT}  ({OUT.stat().st_size / 1024:.0f} KB)")
